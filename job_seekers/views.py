import io
import os
from datetime import date

from bs4 import BeautifulSoup
from django.db.models import Exists, OuterRef
from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from rest_framework import (
    generics, response, status,
    permissions, serializers, filters
)

from core.pagination import CustomPagination
from jobs.models import JobDetails, JobSubCategory, JobCategory
from koor.config.common import Common
from user_profile.models import JobSeekerProfile
from users.models import User
from .models import (
    EducationRecord, EmploymentRecord, JobSeekerLanguageProficiency,
    JobSeekerSkill, AppliedJob, SavedJob, JobPreferences
)
from .serializers import (
    UpdateAboutSerializers, EducationSerializers, JobSeekerLanguageProficiencySerializers,
    EmploymentRecordSerializers, JobSeekerSkillSerializers, AppliedJobSerializers,
    GetAppliedJobsSerializers, GetSavedJobsSerializers, SavedJobSerializers,
    UpdateJobPreferencesSerializers, AdditionalParameterSerializers,
    CategoriesSerializers, ModifyCategoriesSerializers, UpdateAppliedJobSerializers
)


class UpdateAboutView(generics.GenericAPIView):
    """
    A view for updating the JobSeekerProfile of the currently authenticated User.

    Attributes:
        serializer_class: The serializer class to use for updating the JobSeekerProfile.
        permission_classes: The permission classes required to access this view.

    Methods:
        patch: Handle PATCH requests to update the JobSeekerProfile of the authenticated User.

    Returns:
        A Response object with a success or error message, and an appropriate status code.
    """

    serializer_class = UpdateAboutSerializers
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request):
        context = dict()
        if self.request.user.role == "job_seeker":
            profile_instance = get_object_or_404(JobSeekerProfile, user=request.user)
            serializer = self.serializer_class(data=request.data, instance=profile_instance, partial=True)
            try:
                serializer.is_valid(raise_exception=True)
                if 'email' in serializer.validated_data:
                    if User.objects.filter(email__iexact=serializer.validated_data['email']).exists():
                        if profile_instance.user.email__iexact != serializer.validated_data['email']:
                            context['email'] = ["email already in use."]
                            return response.Response(
                                data=context,
                                status=status.HTTP_400_BAD_REQUEST
                            )
                if 'mobile_number' in serializer.validated_data:
                    if User.objects.filter(mobile_number=serializer.validated_data['mobile_number']).exists():
                        if profile_instance.user.mobile_number != serializer.validated_data['mobile_number']:
                            context['mobile_number'] = ["mobile number already in use."]
                            return response.Response(
                                data=context,
                                status=status.HTTP_400_BAD_REQUEST
                            )
                if serializer.update(profile_instance, serializer.validated_data):
                    context['message'] = "Updated Successfully"
                    return response.Response(
                        data=context,
                        status=status.HTTP_200_OK
                    )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )


class EducationsView(generics.GenericAPIView):
    """
    A generic API view for handling education records.

    Attributes:
        - `permission_classes (list)`: The list of permission classes that the user must have to access this view.
                                In this case, only authenticated users can access this view.

        - `serializer_class (EducationSerializers)`: The serializer class that the view uses to serialize and
                                                deserialize the EducationRecord model.

    Returns:
        Serialized education records in JSON format, with authentication required to access the view. 
    """

    permission_classes = [permissions.IsAuthenticated]
    serializer_class = EducationSerializers

    def post(self, request):
        """
        Handles HTTP POST requests for creating a new education record.

        Args:
            request (HttpRequest): The request object sent to the server.

        Returns:
            HTTP response with serialized data in JSON format, and status codes indicating whether the request was
            successful or not.

        Raises:
            serializers.ValidationError: If the data in the request is invalid or incomplete.

            Exception: If an error occurs while processing the request.

        Notes:
            This function creates a new education record using the serializer class specified in the `serializer_class`
            attribute of the view. The serializer is first validated and then saved to the database. The `user` field
            of the education record is set to the currently logged-in user.
            If the serializer is not valid, a `serializers.ValidationError` is raised and a response with a status code
            of 400 is returned. If any other error occurs, a response with a status code of 400 is returned along with
            an error message in the `message` field of the response data. If the serializer is valid and the record is
            successfully created, a response with the serialized data and a status code of 201 is returned.
        """

        context = dict()
        if request.user.role == "job_seeker":
            serializer = self.serializer_class(data=request.data)
            try:
                serializer.is_valid(raise_exception=True)
                serializer.save(user=request.user)
                context["data"] = serializer.data
                return response.Response(
                    data=context,
                    status=status.HTTP_201_CREATED
                )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
            except Exception as e:
                context['message'] = str(e)
                return response.Response(
                    data=context,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def patch(self, request, educationId):
        """
        Update an `EducationRecord` instance for a `job seeker`.

        Args:
            - `request`: The HTTP request object.
            - `educationId`: The ID of the `EducationRecord` instance to update.

        Returns:
            - A Response object with a message and HTTP status code indicating the result of the update.

        Raises:
            - `ValidationError`: If the request data is not valid.
            - `NotFound`: If the EducationRecord instance does not exist.
            - `Exception`: If an error occurs during the update.

        Note:
            - This function is only accessible to job seekers. If the user has a role other than `'job_seeker'`, an
            error message is returned indicating that the user does not have permission to perform the action.
        """
        context = dict()
        if request.user.role == "job_seeker":
            try:
                education_instance = EducationRecord.objects.get(id=educationId, user=request.user)
                serializer = self.serializer_class(data=request.data, instance=education_instance, partial=True)
                try:
                    serializer.is_valid(raise_exception=True)
                    if serializer.update(education_instance, serializer.validated_data):
                        context['message'] = "Updated Successfully"
                        return response.Response(
                            data=context,
                            status=status.HTTP_200_OK
                        )
                except serializers.ValidationError:
                    return response.Response(
                        data=serializer.errors,
                        status=status.HTTP_400_BAD_REQUEST
                    )
            except EducationRecord.DoesNotExist:
                return response.Response(
                    data={"education record": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def delete(self, request, educationId):
        """
        Deletes an EducationRecord object with the given ID if the authenticated user is a job seeker and owns the
        EducationRecord.
        Args:
            request: A DRF request object.
            educationId: An integer representing the ID of the EducationRecord to be deleted.
        Returns:
            A DRF response object with a success or error message and appropriate status code.
        """
        context = dict()
        if request.user.role == "job_seeker":
            try:
                EducationRecord.objects.get(id=educationId, user=request.user).delete()
                context['message'] = "Deleted Successfully"
                return response.Response(
                    data=context,
                    status=status.HTTP_200_OK
                )
            except EducationRecord.DoesNotExist:
                return response.Response(
                    data={"education record": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )


class LanguageView(generics.GenericAPIView):
    """
    A generic API view for handling JobSeekerLanguageProficiency.

    Attributes:
        - `permission_classes (list)`: The list of permission classes that the user must have to access this view.
                                In this case, only authenticated users can access this view.

        - `serializer_class (JobSeekerLanguageProficiencySerializers)`: The serializer class that the view uses to
                                                serialize and deserialize the JobSeekerLanguageProficiency model.

    Returns:
        Serialized JobSeekerLanguageProficiencys in JSON format, with authentication required to access the view.
    """

    permission_classes = [permissions.IsAuthenticated]
    serializer_class = JobSeekerLanguageProficiencySerializers

    def post(self, request):
        """
        Handles HTTP POST requests for creating a new JobSeekerLanguageProficiency.

        Args:
            request (HttpRequest): The request object sent to the server.

        Returns:
            HTTP response with serialized data in JSON format, and status codes indicating whether the request was
            successful or not.

        Raises:
            serializers.ValidationError: If the data in the request is invalid or incomplete.

            Exception: If an error occurs while processing the request.

        Notes:
            This function creates a new JobSeekerLanguageProficiency using the serializer class specified in the `serializer_class`
            attribute of the view. The serializer is first validated and then saved to the database. The `user` field
            of the JobSeekerLanguageProficiency is set to the currently logged-in user.
            If the serializer is not valid, a `serializers.ValidationError` is raised and a response with a status code
            of 400 is returned. If any other error occurs, a response with a status code of 400 is returned along with
            an error message in the `message` field of the response data. If the serializer is valid and the record is
            successfully created, a response with the serialized data and a status code of 201 is returned.
        """

        context = dict()
        if request.user.role == "job_seeker":
            serializer = self.serializer_class(data=request.data)
            try:
                serializer.is_valid(raise_exception=True)
                try:
                    if JobSeekerLanguageProficiency.objects.get(language__title=serializer.validated_data['language'],
                                                                user=request.user):
                        context['language'] = 'Language already in use.'
                        return response.Response(
                            data=context,
                            status=status.HTTP_400_BAD_REQUEST
                        )
                except JobSeekerLanguageProficiency.DoesNotExist:
                    serializer.save(user=request.user)
                    context["data"] = serializer.data
                    return response.Response(
                        data=context,
                        status=status.HTTP_201_CREATED
                    )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
            except Exception as e:
                context['message'] = str(e)
                return response.Response(
                    data=context,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def patch(self, request, languageId):
        """
        Update an `JobSeekerLanguageProficiency` instance for a `job seeker`.

        Args:
            - `request`: The HTTP request object.
            - `languageId`: The ID of the `JobSeekerLanguageProficiency` instance to update.

        Returns:
            - A Response object with a message and HTTP status code indicating the result of the update.

        Raises:
            - `ValidationError`: If the request data is not valid.
            - `NotFound`: If the JobSeekerLanguageProficiency instance does not exist.
            - `Exception`: If an error occurs during the update.

        Note:
            - This function is only accessible to job seekers. If the user has a role other than `'job_seeker'`, an
            error message is returned indicating that the user does not have permission to perform the action.
        """
        context = dict()
        if request.user.role == "job_seeker":
            try:
                language_instance = JobSeekerLanguageProficiency.objects.get(id=languageId, user=request.user)
                serializer = self.serializer_class(data=request.data, instance=language_instance, partial=True)
                try:
                    serializer.is_valid(raise_exception=True)
                    if serializer.update(language_instance, serializer.validated_data):
                        context['message'] = "Updated Successfully"
                        return response.Response(
                            data=context,
                            status=status.HTTP_200_OK
                        )
                except serializers.ValidationError:
                    return response.Response(
                        data=serializer.errors,
                        status=status.HTTP_400_BAD_REQUEST
                    )
            except JobSeekerLanguageProficiency.DoesNotExist:
                return response.Response(
                    data={"job_seeker_language_proficiency": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def delete(self, request, languageId):
        """
        Deletes an JobSeekerLanguageProficiency object with the given ID if the authenticated user is a
        job seeker and owns the JobSeekerLanguageProficiency.
        Args:
            request: A DRF request object.
            languageId: An integer representing the ID of the JobSeekerLanguageProficiency to be deleted.
        Returns:
            A DRF response object with a success or error message and appropriate status code.
        """
        context = dict()
        if request.user.role == "job_seeker":
            try:
                JobSeekerLanguageProficiency.objects.get(id=languageId, user=request.user).delete()
                context['message'] = "Deleted Successfully"
                return response.Response(
                    data=context,
                    status=status.HTTP_200_OK
                )
            except JobSeekerLanguageProficiency.DoesNotExist:
                return response.Response(
                    data={"job_seeker_language_proficiency": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )


class WorkExperiencesView(generics.GenericAPIView):
    """
    A generic API view for handling employment record.

    Attributes:
        - `permission_classes (list)`: The list of permission classes that the user must have to access this view.
                                In this case, only authenticated users can access this view.

        - `serializer_class (EmploymentRecordSerializers)`: The serializer class that the view uses to serialize and
                                                deserialize the EmploymentRecord model.

    Returns:
        Serialized employment record in JSON format, with authentication required to access the view. 
    """

    permission_classes = [permissions.IsAuthenticated]
    serializer_class = EmploymentRecordSerializers

    def post(self, request):
        """
        Handles HTTP POST requests for creating a new employment record.

        Args:
            request (HttpRequest): The request object sent to the server.

        Returns:
            HTTP response with serialized data in JSON format, and status codes indicating whether the request was
            successful or not.

        Raises:
            serializers.ValidationError: If the data in the request is invalid or incomplete.

            Exception: If an error occurs while processing the request.

        Notes:
            This function creates a new employment record using the serializer class specified in the `serializer_class`
            attribute of the view. The serializer is first validated and then saved to the database. The `user` field
            of the employment record is set to the currently logged-in user.
            If the serializer is not valid, a `serializers.ValidationError` is raised and a response with a status code
            of 400 is returned. If any other error occurs, a response with a status code of 400 is returned along with
            an error message in the `message` field of the response data. If the serializer is valid and the record is
            successfully created, a response with the serialized data and a status code of 201 is returned.
        """

        context = dict()
        if request.user.role == "job_seeker":
            serializer = self.serializer_class(data=request.data)
            try:
                serializer.is_valid(raise_exception=True)
                serializer.save(user=request.user)
                context["data"] = serializer.data
                return response.Response(
                    data=context,
                    status=status.HTTP_201_CREATED
                )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
            except Exception as e:
                context['message'] = str(e)
                return response.Response(
                    data=context,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def patch(self, request, workExperienceId):
        """
        Update an `EmploymentRecord` instance for a `job seeker`.

        Args:
            - `request`: The HTTP request object.
            - `workExperienceId`: The ID of the `EmploymentRecord` instance to update.

        Returns:
            - A Response object with a message and HTTP status code indicating the result of the update.

        Raises:
            - `ValidationError`: If the request data is not valid.
            - `NotFound`: If the EmploymentRecord instance does not exist.
            - `Exception`: If an error occurs during the update.

        Note:
            - This function is only accessible to job seekers. If the user has a role other than `'job_seeker'`, an
            error message is returned indicating that the user does not have permission to perform the action.
        """
        context = dict()
        if request.user.role == "job_seeker":
            try:
                employment_instance = EmploymentRecord.objects.get(id=workExperienceId, user=request.user)
                serializer = self.serializer_class(data=request.data, instance=employment_instance, partial=True)
                try:
                    serializer.is_valid(raise_exception=True)
                    if serializer.update(employment_instance, serializer.validated_data):
                        context['message'] = "Updated Successfully"
                        return response.Response(
                            data=context,
                            status=status.HTTP_200_OK
                        )
                except serializers.ValidationError:
                    return response.Response(
                        data=serializer.errors,
                        status=status.HTTP_400_BAD_REQUEST
                    )
            except EmploymentRecord.DoesNotExist:
                return response.Response(
                    data={"employment_record": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def delete(self, request, workExperienceId):
        """
        Deletes an EmploymentRecord object with the given ID if the authenticated user is a job seeker and owns the
        EmploymentRecord.
        Args:
            request: A DRF request object.
            workExperienceId: An integer representing the ID of the EmploymentRecord to be deleted.
        Returns:
            A DRF response object with a success or error message and appropriate status code.
        """
        context = dict()
        if request.user.role == "job_seeker":
            try:
                EmploymentRecord.objects.get(id=workExperienceId, user=request.user).delete()
                context['message'] = "Deleted Successfully"
                return response.Response(
                    data=context,
                    status=status.HTTP_200_OK
                )
            except EmploymentRecord.DoesNotExist:
                return response.Response(
                    data={"employment_record": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )


class SkillsView(generics.GenericAPIView):
    """
    A generic API view for handling skills.

    Attributes:
        - `permission_classes (list)`: The list of permission classes that the user must have to access this view.
                                In this case, only authenticated users can access this view.

        - `serializer_class (JobSeekerSkillSerializers)`: The serializer class that the view uses to serialize and
                                                deserialize the JobSeekerSkill model.

    Returns:
        Serialized skills in JSON format, with authentication required to access the view. 
    """

    permission_classes = [permissions.IsAuthenticated]
    serializer_class = JobSeekerSkillSerializers

    def post(self, request):
        """
        Handles HTTP POST requests for creating a new skills.

        Args:
            request (HttpRequest): The request object sent to the server.

        Returns:
            HTTP response with serialized data in JSON format, and status codes indicating whether the request was
            successful or not.

        Raises:
            serializers.ValidationError: If the data in the request is invalid or incomplete.

            Exception: If an error occurs while processing the request.

        Notes:
            This function creates a new skills using the serializer class specified in the `serializer_class`
            attribute of the view. The serializer is first validated and then saved to the database. The `user` field
            of the skills is set to the currently logged-in user.
            If the serializer is not valid, a `serializers.ValidationError` is raised and a response with a status code
            of 400 is returned. If any other error occurs, a response with a status code of 400 is returned along with
            an error message in the `message` field of the response data. If the serializer is valid and the record is
            successfully created, a response with the serialized data and a status code of 201 is returned.
        """

        context = dict()
        if request.user.role == "job_seeker":
            serializer = self.serializer_class(data=request.data)
            try:
                serializer.is_valid(raise_exception=True)
                for data in serializer.validated_data:
                    try:
                        if JobSeekerSkill.objects.get(skill_id=data, user=request.user):
                            pass
                    except JobSeekerSkill.DoesNotExist:
                        JobSeekerSkill.objects.create(skill_id=data, user=request.user)
                context["message"] = "Skills added."
                return response.Response(
                    data=context,
                    status=status.HTTP_201_CREATED
                )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
            except Exception as e:
                context['message'] = str(e)
                return response.Response(
                    data=context,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )


class JobsApplyView(generics.ListAPIView):
    """
    A view for retrieving a list of applied jobs.

    This view supports HTTP GET requests and returns a list of applied jobs for the authenticated user.
    The applied jobs are serialized using the `GetAppliedJobsSerializers` class.

    This view requires the user to be authenticated, and uses the `IsAuthenticated` permission class.
    The view supports searching the applied jobs by job title, using the `SearchFilter` filter backend.

    Attributes:
        - `serializer_class`: The serializer class to use for serializing the applied jobs.
        - `permission_classes`: A list of permission classes that the user must pass in order to access this view.
        - `queryset`: The base queryset for the view. This attribute is not used in this view, since the queryset
            is dynamically generated in the `get_queryset` method.
        - `filter_backends`: A list of filter backends to use for filtering the applied jobs.
        - `search_fields`: The fields to search for when filtering the applied jobs.
    """

    serializer_class = GetAppliedJobsSerializers
    permission_classes = [permissions.IsAuthenticated]
    queryset = None
    filter_backends = [filters.SearchFilter]
    search_fields = ['job__title']
    pagination_class = CustomPagination

    def list(self, request):
        """
        Returns a paginated list of serialized applied jobs for the authenticated user.

        This method returns a paginated list of applied jobs for the authenticated user. The applied jobs are
        serialized using the `GetAppliedJobsSerializers` class.

        Args:
            request: The HTTP request object.

        Returns:
            A HTTP response object containing a paginated list of serialized applied jobs.

        The response includes the following fields:
            - `count (int)`: The total number of applied jobs for the authenticated user.
            - `next (str)`: The URL for the next page of results, or null if there are no more pages.
            - `previous (str)`: The URL for the previous page of results, or null if this is the first page.
            - `results (list)`: A list of serialized applied jobs for the authenticated user.

        """

        queryset = self.filter_queryset(self.get_queryset())
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True, context={"request": request})
            return self.get_paginated_response(serializer.data)
        serializer = self.get_serializer(queryset, many=True, context={"request": request})
        return response.Response(serializer.data)

    def post(self, request, jobId):
        """
        Creates a new application for a job posting by a job seeker.

        Args:
            request: The HTTP request object.
            jobId (int): The ID of the job posting to apply for.

        Returns:
            A response object with the following keys:
            - "message" (str): A message indicating the success or failure of the request.

            - If the request is successful, the response will have a status code of 200 (HTTP_200_OK).
            - If the user does not have permission to perform this action, the response will have a status
            code of 401 (HTTP_401_UNAUTHORIZED).
            - If the specified job posting does not exist, the response will have a status code of 404
            (HTTP_404_NOT_FOUND).
            - If there is an error while processing the request, the response will have a status code of 404
            (HTTP_404_NOT_FOUND) and a message describing the error.
        """

        context = dict()
        if request.user.role == "job_seeker":
            try:
                job_instance = JobDetails.objects.get(id=jobId)
                try:
                    if AppliedJob.objects.get(job=job_instance, user=request.user):
                        context["message"] = "You are already applied"
                        return response.Response(
                            data=context,
                            status=status.HTTP_400_BAD_REQUEST
                        )
                except AppliedJob.DoesNotExist:
                    if request.user.name and request.user.user_profile_jobseekerprofile_user.gender and request.user.user_profile_jobseekerprofile_user.dob and request.user.user_profile_jobseekerprofile_user.employment_status and request.user.user_profile_jobseekerprofile_user.country and request.user.user_profile_jobseekerprofile_user.city and request.user.user_profile_jobseekerprofile_user.highest_education:
                        serializer = AppliedJobSerializers(data=request.data)
                        try:
                            serializer.is_valid(raise_exception=True)
                            serializer.save(user=request.user, job_instance=job_instance)
                            context["message"] = "Applied Successfully"
                            return response.Response(
                                data=context,
                                status=status.HTTP_200_OK
                            )
                        except serializers.ValidationError:
                            return response.Response(
                                data=str(serializer.errors),
                                status=status.HTTP_400_BAD_REQUEST
                            )
                    else:
                        return response.Response(
                            data={"message": ["Please complete your profile."]},
                            status=status.HTTP_404_NOT_FOUND
                        )
            except JobDetails.DoesNotExist:
                return response.Response(
                    data={"job": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = str(e)
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )

        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def put(self, request, jobId):
        """
        Update an existing job application for a job seeker user.

        Args:
            - `request`: An HTTP request object.
            - `jobId (int)`: The ID of the job application to update.

        Returns:
            A Response object with an appropriate HTTP status code and message.

        Raises:
            - `ValidationError`: If the request data is invalid.
            - `Http404`: If the specified job or job application does not exist.
            - `Exception`: If any other error occurs.

        The function first checks if the user making the request is a job seeker. If not, it returns an
        unauthorized response. Otherwise, it tries to update the specified job application with the given
        data using an instance of the UpdateJobSerializers class. If the update is successful, it returns
        a success response with a message indicating that the update was successful. If the job application
        or job do not exist, it returns a 404 response with an appropriate error message. If the request data
        is invalid or if any other error occurs, it returns a 400 or 404 response with an error message
        containing details about the error.
        """

        context = dict()
        if self.request.user.role == "job_seeker":
            serializer = UpdateAppliedJobSerializers(data=request.data)
            try:
                job_instance = JobDetails.objects.get(id=jobId)
                try:
                    applied_job = AppliedJob.objects.get(job=job_instance, user=request.user)
                    if applied_job.shortlisted_at or applied_job.rejected_at or applied_job.created.date() < date.today():
                        context['message'] = "You cannot update this applied job"
                        return response.Response(
                            data=context,
                            status=status.HTTP_200_OK
                        )
                    else:
                        serializer.is_valid(raise_exception=True)
                        if serializer.update(applied_job, serializer.validated_data):
                            context['message'] = "Updated Successfully"
                            return response.Response(
                                data=context,
                                status=status.HTTP_200_OK
                            )
                except serializers.ValidationError:
                    return response.Response(
                        data=serializer.errors,
                        status=status.HTTP_400_BAD_REQUEST
                    )
                except AppliedJob.DoesNotExist:
                    return response.Response(
                        data={"application": "Does Not Exist"},
                        status=status.HTTP_404_NOT_FOUND
                    )
            except JobDetails.DoesNotExist:
                return response.Response(
                    data={"job": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = str(e)
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def delete(self, request, jobId):
        """
        Deletes an AppliedJob object with the given job if the authenticated user is a job seeker and owns the
        AppliedJob.
        Args:
            request: A DRF request object.
            educationId: An integer representing the ID of the AppliedJob to be deleted.
        Returns:
            A DRF response object with a success or error message and appropriate status code.
        """
        context = dict()
        if request.user.role == "job_seeker":
            # print(timezone.now())
            try:
                job_instance = JobDetails.objects.get(id=jobId)
                try:
                    applied_job = AppliedJob.all_objects.get(job=job_instance, user=request.user)
                    if applied_job.shortlisted_at or applied_job.rejected_at or applied_job.created.date() < date.today():
                        context['message'] = "You cannot revoke this applied job"
                    else:
                        applied_job.delete(soft=False)
                        context['message'] = "Revoked applied job"
                    return response.Response(
                        data=context,
                        status=status.HTTP_200_OK
                    )
                except AppliedJob.DoesNotExist:
                    return response.Response(
                        data={"AppliedJob": "Does Not Exist"},
                        status=status.HTTP_404_NOT_FOUND
                    )
            except JobDetails.DoesNotExist:
                return response.Response(
                    data={"job": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def get_queryset(self, **kwargs):
        """
        Returns the queryset of applied jobs for the authenticated user.

        This method returns a queryset of AppliedJob objects for the authenticated user, ordered by their creation date
        in descending order.

        Args:
            **kwargs: Additional keyword arguments.

        Returns:
            A queryset of AppliedJob objects for the authenticated user, ordered by their creation date in descending
            order.
        """
        order_by = None
        if 'search_by' in self.request.GET:
            search_by = self.request.GET['search_by']
            if search_by == 'salary':
                order_by = 'job__budget_amount'
            elif search_by == 'expiration':
                order_by = 'job__deadline'
            elif search_by == 'created_at':
                order_by = 'job__created'
            if order_by:
                if 'order_by' in self.request.GET:
                    if 'descending' in self.request.GET['order_by']:
                        return AppliedJob.objects.filter(
                            user=self.request.user,
                            job__is_removed=False
                        ).order_by("-" + str(order_by))
                    else:
                        return AppliedJob.objects.filter(
                            user=self.request.user,
                            job__is_removed=False
                        ).order_by(str(order_by))
                else:
                    return AppliedJob.objects.filter(
                        user=self.request.user,
                        job__is_removed=False
                    ).order_by(str(order_by))
        return AppliedJob.objects.filter(
            user=self.request.user,
            job__is_removed=False
        )


class JobsSaveView(generics.ListAPIView):
    """
    A view for retrieving a list of saved jobs.

    This view supports HTTP GET requests and returns a list of saved jobs for the authenticated user.
    The saved jobs are serialized using the `GetSavedJobsSerializers` class.

    This view requires the user to be authenticated, and uses the `IsAuthenticated` permission class.
    The view supports searching the saved jobs by job title, using the `SearchFilter` filter backend.

    Attributes:
        - `serializer_class`: The serializer class to use for serializing the saved jobs.
        - `permission_classes`: A list of permission classes that the user must pass in order to access this view.
        - `queryset`: The base queryset for the view. This attribute is not used in this view, since the queryset
            is dynamically generated in the `get_queryset` method.
        - `filter_backends`: A list of filter backends to use for filtering the saved jobs.
        - `search_fields`: The fields to search for when filtering the saved jobs.
    """
    serializer_class = GetSavedJobsSerializers
    permission_classes = [permissions.IsAuthenticated]
    queryset = None
    filter_backends = [filters.SearchFilter]
    search_fields = ['title']
    pagination_class = CustomPagination

    def list(self, request):
        """
        Returns a paginated list of serialized saved jobs for the authenticated user.

        This method returns a paginated list of saved jobs for the authenticated user. The saved jobs are
        serialized using the `GetSavedJobsSerializers` class.

        Args:
            request: The HTTP request object.

        Returns:
            A HTTP response object containing a paginated list of serialized saved jobs.

        The response includes the following fields:
            - `count (int)`: The total number of saved jobs for the authenticated user.
            - `next (str)`: The URL for the next page of results, or null if there are no more pages.
            - `previous (str)`: The URL for the previous page of results, or null if this is the first page.
            - `results (list)`: A list of serialized saved jobs for the authenticated user.

        """

        queryset = self.filter_queryset(self.get_queryset())
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True, context={"request": request})
            return self.get_paginated_response(serializer.data)
        serializer = self.get_serializer(queryset, many=True, context={"request": request})
        return response.Response(serializer.data)

    def post(self, request, jobId):
        """
        Creates a new application for a job posting by a job seeker.

        Args:
            request: The HTTP request object.
            jobId (int): The ID of the job posting to apply for.

        Returns:
            A response object with the following keys:
            - "message" (str): A message indicating the success or failure of the request.

            - If the request is successful, the response will have a status code of 200 (HTTP_200_OK).
            - If the user does not have permission to perform this action, the response will have a status
            code of 401 (HTTP_401_UNAUTHORIZED).
            - If the specified job posting does not exist, the response will have a status code of 404
            (HTTP_404_NOT_FOUND).
            - If there is an error while processing the request, the response will have a status code of 404
            (HTTP_404_NOT_FOUND) and a message describing the error.
        """

        context = dict()
        if request.user.role == "job_seeker":
            try:
                job_instance = JobDetails.objects.get(id=jobId)
                try:
                    if SavedJob.objects.get(job=job_instance, user=request.user):
                        context["message"] = "You are already saved"
                        return response.Response(
                            data=context,
                            status=status.HTTP_400_BAD_REQUEST
                        )
                except SavedJob.DoesNotExist:
                    serializer = SavedJobSerializers(data=request.data)
                    try:
                        serializer.is_valid(raise_exception=True)
                        serializer.save(user=request.user, job_instance=job_instance)
                        context["message"] = "Saved Successfully"
                        return response.Response(
                            data=context,
                            status=status.HTTP_200_OK
                        )
                    except serializers.ValidationError:
                        return response.Response(
                            data=str(serializer.errors),
                            status=status.HTTP_400_BAD_REQUEST
                        )
            except JobDetails.DoesNotExist:
                return response.Response(
                    data={"job": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = str(e)
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )

        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def delete(self, request, jobId):
        """
        Deletes an SavedJob object with the given job if the authenticated user is a job seeker and owns the
        SavedJob.
        Args:
            request: A DRF request object.
            educationId: An integer representing the ID of the SavedJob to be deleted.
        Returns:
            A DRF response object with a success or error message and appropriate status code.
        """
        context = dict()
        if request.user.role == "job_seeker":
            try:
                job_instance = JobDetails.objects.get(id=jobId)
                try:
                    SavedJob.all_objects.get(job=job_instance, user=request.user).delete(soft=False)
                    context['message'] = "Job Unsaved"
                    return response.Response(
                        data=context,
                        status=status.HTTP_200_OK
                    )
                except SavedJob.DoesNotExist:
                    return response.Response(
                        data={"savedJobId": "Does Not Exist"},
                        status=status.HTTP_404_NOT_FOUND
                    )
            except JobDetails.DoesNotExist:
                return response.Response(
                    data={"job": "Does Not Exist"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                context["message"] = e
                return response.Response(
                    data=context,
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def get_queryset(self, **kwargs):
        """
        Returns the queryset of saved jobs for the authenticated user.

        This method returns a queryset of SavedJob objects for the authenticated user, ordered by their creation date
        in descending order.

        Args:
            **kwargs: Additional keyword arguments.

        Returns:
            A queryset of SavedJob objects for the authenticated user, ordered by their creation date in descending
            order.
        """
        order_by = None
        if 'search_by' in self.request.GET:
            search_by = self.request.GET['search_by']
            if search_by == 'salary':
                order_by = 'job__budget_amount'
            elif search_by == 'expiration':
                order_by = 'job__deadline'
            elif search_by == 'created_at':
                order_by = 'job__created'
            if order_by:
                if 'order_by' in self.request.GET:
                    if 'descending' in self.request.GET['order_by']:
                        return SavedJob.objects.filter(
                            user=self.request.user,
                            job__is_removed=False
                        ).order_by("-" + str(order_by))
                    else:
                        return SavedJob.objects.filter(
                            user=self.request.user,
                            job__is_removed=False
                        ).order_by(str(order_by))
                else:
                    return SavedJob.objects.filter(
                        user=self.request.user,
                        job__is_removed=False
                    ).order_by(str(order_by))
        return SavedJob.objects.filter(
            user=self.request.user,
            job__is_removed=False
        )


class UpdateJobPreferencesView(generics.GenericAPIView):
    """
        A view class for updating `JobPreferences` instances via PATCH requests.

        This class extends the GenericAPIView class provided by Django REST framework, and specifies the serializer
        class and permission classes to use.

        Attributes:
            - `serializer_class`: The serializer class used to serialize and validate request data.
            - `permission_classes`: A list of permission classes to use for authorization.
    """

    serializer_class = UpdateJobPreferencesSerializers
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request):
        """
            Handle a PATCH request to update a JobPreferences instance.

            This method checks that the requesting user is a job seeker, and retrieves or creates a `JobPreferences`
            instance for the user. It then uses the serializer class to validate and update the instance with the
            request data.

            Args:
                - `request`: The `PATCH` request object.

            Returns:
                - A response object with a success or error message, and an HTTP status code.
        """

        context = dict()
        if self.request.user.role == "job_seeker":
            if JobPreferences.objects.filter(user=request.user).exists():
                preference_instance = get_object_or_404(JobPreferences, user=request.user)
            else:
                preference_instance = JobPreferences.objects.create(user=request.user)

            serializer = self.serializer_class(data=request.data, instance=preference_instance, partial=True)
            try:
                serializer.is_valid(raise_exception=True)
                if serializer.update(preference_instance, serializer.validated_data):
                    context['message'] = "Updated Successfully"
                    return response.Response(
                        data=context,
                        status=status.HTTP_200_OK
                    )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )


class AdditionalParameterView(generics.GenericAPIView):
    """
    A view for handling PUT requests to update JobSeekerProfile instances with additional boolean fields.

    Attributes:
        - `serializer_class (AdditionalParameterSerializers)`: The serializer class to use for serializing and
            deserializing JobSeekerProfile instances with additional boolean fields.
        - `permission_classes ([permissions.IsAuthenticated])`: The permission classes to use for this view, which
            require authentication for access.

    Methods:
        - `put(request)`: Handles PUT requests to update a JobSeekerProfile instance with additional boolean fields.

    Returns:
        - `response.Response`: A response object containing a success or error message and an HTTP status code.
    """

    serializer_class = AdditionalParameterSerializers
    permission_classes = [permissions.IsAuthenticated]

    def put(self, request):
        context = dict()
        if self.request.user.role == "job_seeker":
            profile_instance = get_object_or_404(JobSeekerProfile, user=request.user)
            serializer = self.serializer_class(data=request.data, instance=profile_instance, partial=True)
            try:
                serializer.is_valid(raise_exception=True)
                if serializer.update(profile_instance, serializer.validated_data):
                    context['message'] = "Updated Successfully"
                    return response.Response(
                        data=context,
                        status=status.HTTP_200_OK
                    )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=context,
                status=status.HTTP_401_UNAUTHORIZED
            )


class CategoryView(generics.GenericAPIView):
    """
    A view that returns a list of job seeker categories and their associated sub-categories.

    Attributes:
        - `serializer_class (CategoriesSerializers)`: The serializer class used for serializing the data.
        - `permission_classes (list)`: The permission classes required for accessing the view.

    Methods:
        - `get`: Retrieves a list of job seeker categories and their associated sub-categories.
    """

    serializer_class = ModifyCategoriesSerializers
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        """
        Retrieve a list of job seeker categories and their associated sub-categories.

        If the user is a job seeker, the method returns a list of categories and their associated sub-categories.
        Otherwise, it returns an HTTP 401 Unauthorized response with an error message.

        Args:
            - `request (Request)`: The HTTP request object.

        Returns:
            - A Response object containing the serialized data and an HTTP status code.

        Raises:
            None.
        """

        response_context = dict()
        if self.request.user.role == "job_seeker":
            category_data = JobCategory.objects.filter(is_removed=False).annotate(
                has_subcategory=Exists(JobSubCategory.objects.filter(category_id=OuterRef('id')))
            ).filter(has_subcategory=True)
            get_data = CategoriesSerializers(category_data, many=True, context={'user': request.user})
            return response.Response(
                data=get_data.data,
                status=status.HTTP_200_OK
            )
        else:
            response_context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=response_context,
                status=status.HTTP_401_UNAUTHORIZED
            )

    def put(self, request):
        """
        Handle HTTP `PUT` requests.

        This method is used to update an existing user profile for job seekers only.
        If the user is not a job seeker, a 401 UNAUTHORIZED response is returned.

        Args:
            - `request (HttpRequest)`: The HTTP request object containing the user profile data to update.

        Returns:
            A Response object with either a `200 OK` status code and a success message if the `update was successful`,
            or a `400 BAD REQUEST` status code and a dictionary of `error` messages if the data was
            `invalid or incomplete`.
            If the user is not a job seeker, a `401 UNAUTHORIZED` status code and an error message are returned.
        """

        response_context = dict()
        if self.request.user.role == "job_seeker":
            serializer = self.serializer_class(data=request.data)
            try:
                serializer.is_valid(raise_exception=True)
                serializer.save(user=request.user)
                response_context['message'] = "Updated Successfully"
                return response.Response(
                    data=response_context,
                    status=status.HTTP_200_OK
                )
            except serializers.ValidationError:
                return response.Response(
                    data=serializer.errors,
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            response_context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=response_context,
                status=status.HTTP_401_UNAUTHORIZED
            )


def RemoveAvailability():
    """
    Sets the availability flag of all JobPreferences objects to False.

    Returns:
        - `HttpResponse`: A response indicating that the operation has been completed.
    """
    JobPreferences.objects.all().update(is_available=False)
    return HttpResponse("done")


class ResumeView(generics.GenericAPIView):
    """
    View for generating a resume in Microsoft Word format (.docx).

    The view supports only authenticated users with a role of 'job_seeker'.
    The generated resume includes the user's name, contact information, profile summary, skills, education,
    employment summary, and language proficiency.

    The generated file is saved to the media folder and the response includes the URL to download the file.

    Note:
        This view uses the `python-docx` library for generating the Word document.

    Serializer:
        ModifyCategoriesSerializers

    Permissions:
        IsAuthenticated

    Methods:
        get -- Generates a resume based on the logged-in user's information.

    Return:
        A JSON response containing the URL to download the generated file.

    Example:
        {
            "url": "/media/123.docx"
        }
    """

    serializer_class = ModifyCategoriesSerializers
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):

        response_context = dict()
        if self.request.user.role == "job_seeker":
            profile_instance = get_object_or_404(JobSeekerProfile, user=request.user)
            downloaded_file = download_word_document(request, profile_instance)
            return response.Response(
                data={"url": "/media/" + downloaded_file},
                status=status.HTTP_200_OK
            )
        else:
            response_context['message'] = "You do not have permission to perform this action."
            return response.Response(
                data=response_context,
                status=status.HTTP_401_UNAUTHORIZED
            )


def add_horizontal_line(doc):
    # Add a paragraph with a border to mimic the horizontal line
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set the alignment to center
    paragraph.add_run().add_break()  # Add an empty line to create some space above the line
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Set line spacing to single
    border = paragraph.border_bottom
    border.color.rgb = (0, 0, 0)  # Set border color to black
    border.width = Pt(6)  # Set border width to 6pt (adjust as needed)


def get_element_qname(element):
    """
    Get the fully qualified name of the element.
    """
    return f"{{{element.nsmap[element.prefix]}}}{element.tag.split('}')[-1]}"


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to the paragraph with the given URL and display text.
    """
    run = paragraph.add_run("Email : ")
    hyperlink_run = paragraph.add_run()
    hyperlink_run.add_text(text)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(get_element_qname(OxmlElement('w:rId')), url)
    hyperlink_run._r.append(hyperlink)
    hyperlink_run.font.underline = True
    hyperlink_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Set the hyperlink color to blue


def add_content_to_document(element, doc, processed_elements):
    if element in processed_elements:
        return

    processed_elements.add(element)

    if element.name == 'h1':
        # Add the title with 'Title' style
        doc.add_paragraph(element.text, style='Title')

    elif element.name == 'h4':
        # Create a new paragraph for <h4> tag
        paragraph = doc.add_paragraph()
        if element.get('class') == ['align_right']:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Set the alignment to right
            # Set the line spacing to 1.5 (you can change the value to your desired line spacing)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = 0
        elif element.get('class') == ['align_left']:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set the alignment to left
            # Set the line spacing to 1.5 (you can change the value to your desired line spacing)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = 12

        # Set the text and apply formatting
        run = paragraph.add_run(element.text)
        run.bold = True
        font = run.font
        font.size = Pt(14)  # Set the font size to 14pt (adjust as needed)

    elif element.name == 'h6':
        # Create a new paragraph for <h4> tag
        paragraph = doc.add_paragraph()
        # Set the text and apply formatting
        run = paragraph.add_run(element.text)
        run.bold = True
        font = run.font
        font.size = Pt(11)  # Set the font size to 14pt (adjust as needed)

    elif element.name == 'p':
        if element.strong:
            add_content_to_document(element.strong, doc, processed_elements)
        else:
            # Create a new paragraph for <span> tag
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set the alignment to right

            # Set the text and apply formatting
            run = paragraph.add_run(element.text)
            font = run.font
            font.size = Pt(11)  # Set the font size to 11pt (adjust as needed)

    elif element.name == 'strong':
        # Create a new paragraph for <h4> tag
        paragraph = doc.add_paragraph()
        # Set the text and apply formatting
        run = paragraph.add_run(element.text)
        run.bold = True
        font = run.font
        font.size = Pt(11)  # Set the font size to 14pt (adjust as needed)

    elif element.name == 'li':
        
        # Create a new paragraph for <h4> tag
        paragraph = doc.add_paragraph(style='List Bullet')
        
        # Set a custom paragraph indentation to position the text after the bullet
        paragraph.paragraph_format.left_indent = Cm(1.5)

        run = paragraph.add_run(element.text)  # Add the text
        font = run.font
        font.size = Pt(11)  # Set the font size of the text (adjust as needed)

    elif element.name == 'hr':
        # Handle <hr> tag, add a horizontal line
        paragraph = doc.add_paragraph('', style='Intense Quote')
        paragraph.alignment = 1  # This sets the alignment to centered
        # Set the line spacing to 1.5 (you can change the value to your desired line spacing)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = 0
        paragraph_format.space_before = 0
        paragraph_format.line_spacing = 0.6

        # Set the font size to 1 for the content of the paragraph
        run = paragraph.add_run('')  # Add an empty space to the paragraph
        font = run.font
        font.size = Pt(1)  # Pt() is a function to set font size in points

        paragraph.style.paragraph_format.left_indent = -1440  # Set a negative left indent to expand the width
        paragraph.style.paragraph_format.right_indent = -1440  # Set a negative right indent to expand the width

    elif element.name == 'span' and element.get('class') == ['align_right']:
        # Handle <span> with class 'align_right'
        if element.a:
            # If <span> contains an <a> tag, process it separately
            add_content_to_document(element.a, doc, processed_elements)
        else:
            # Create a new paragraph for <span> tag
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Set the alignment to right
            # Set the line spacing to 1.5 (you can change the value to your desired line spacing)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = 0
            # Set the text and apply formatting
            run = paragraph.add_run(element.text)
            font = run.font
            font.size = Pt(11)  # Set the font size to 11pt (adjust as needed)

    elif element.name == 'span' and element.get('class') == ['education']:
        # Handle <span> with class 'align_right'
        if element.a:
            # If <span> contains an <a> tag, process it separately
            add_content_to_document(element.a, doc, processed_elements)
        else:
            # Create a new paragraph for <span> tag
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set the alignment to right
            # Set the line spacing to 1.5 (you can change the value to your desired line spacing)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = 0

            # Set the text and apply formatting
            run = paragraph.add_run(element.text)
            font = run.font
            font.size = Pt(11)  # Set the font size to 11pt (adjust as needed)

    elif element.name == 'a':
        # Handle <a> tag
        url = element['href']
        text = element.text.strip()
        paragraph = doc.add_paragraph(style='Normal')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = 0
        paragraph_format.space_before = 0
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        add_hyperlink(paragraph, url, text)

    elif element.string and element.string.strip():
        # Add other non-empty text elements as regular paragraphs with 'Normal' style
        doc.add_paragraph(element.string.strip(), style='Normal')


    elif element.string and element.string.strip():
        # Add other non-empty text elements as regular paragraphs
        doc.add_paragraph(element.string.strip())

    # Recursively process child elements
    for child in element.children:
        if child.name:
            add_content_to_document(child, doc, processed_elements)


def download_word_document(request, profile_instance):
    sills_data = JobSeekerSkill.objects.filter(user=profile_instance.user)
    education_data = EducationRecord.objects.filter(user=profile_instance.user)
    employment_data = EmploymentRecord.objects.filter(user=profile_instance.user)
    languages_data = JobSeekerLanguageProficiency.objects.filter(user=profile_instance.user)
    mobile_number = profile_instance.user.mobile_number
    new_mobile_number = " "
    for i in range(0, len(mobile_number), 5):
        new_mobile_number += mobile_number[i:i + 5] + " "

    # Sample HTML content to be converted to DOCX
    html_content = render_to_string('resume.html', {
        'profile_instance': profile_instance, 'new_mobile_number': new_mobile_number, 'sills_data': sills_data,
        'education_data': education_data, 'employment_data': employment_data, 'languages_data': languages_data,
    })

    # Create a new Document
    doc = Document()

    # Convert HTML to plain text using BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Process the HTML content and add to the Word document
    processed_elements = set()
    for element in soup.recursiveChildGenerator():
        if element.name:
            add_content_to_document(element, doc, processed_elements)

    # Create a file path for the DOC file in the media folder
    media_root = Common.MEDIA_ROOT
    file_name = str(profile_instance.id) + '.docx'
    file_path = os.path.join(media_root, file_name)

    # Save the document to the file path
    doc.save(file_path)

    # Create a buffer to store the document
    output_buffer = io.BytesIO()
    doc.save(output_buffer)

    # Set the appropriate headers for the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="converted_document.docx"'

    # Write the content of the buffer to the response
    response.write(output_buffer.getvalue())

    return file_name


class ResumeUseridView(generics.GenericAPIView):

    permission_classes = [permissions.AllowAny]

    def get(self, request):

        response_context = dict()
        downloaded_file = ""
        if 'user-id' in request.GET:
            profile_instance = get_object_or_404(JobSeekerProfile, user__id=request.GET['user-id'])
            downloaded_file = download_word_document(request, profile_instance)
            
        return HttpResponseRedirect(Common.BASE_URL + "/media/" + downloaded_file)
